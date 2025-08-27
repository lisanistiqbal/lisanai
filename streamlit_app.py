
import pandas as pd
import base64
import streamlit as st
import json
import time
import re
from lxml import etree
from io import StringIO, BytesIO
import requests 
from streamlit_lottie import st_lottie  
import vertexai
from typing import List
from google.cloud import translate, aiplatform
from vertexai.generative_models import GenerativeModel, Part, SafetySetting
import os
#from pydub import AudioSegment
import google.generativeai as genai
import pyperclip
import PyPDF2
import docx
import xml.etree.ElementTree as ET
import pickle
from pathlib import Path
import streamlit_authenticator as stauth
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P 
import polib
import subprocess
import time
from lxml import etree
import tempfile

# --- USER AUTHENTICATION ---
names = ["Asif Iqbal", "Sadullah Saad", "Faheem Ahmad"]
usernames = ["asif", "saad", "fahmad"]

# Load the hashed passwords from the .pkl file
with open("hashed_passwords.pkl", "rb") as file:
    hashed_passwords = pickle.load(file)  # This should be a dictionary

credentials = {
    "usernames": {
        username: {
            "name": username.replace("_", " ").title(),  # Auto-format names
            "password": hashed_passwords[username]  # Use the hashed password from the .pkl file
        }
        for username in hashed_passwords
    }
}

# load hashed passwords

authenticator = stauth.Authenticate(
    credentials, 
    cookie_name="some_cookie_name",
    key="some_secret_key",
    cookie_expiry_days=30
)


#name, authentication_status, username = authenticator.login('Login', 'sidebar')

#if authentication_status == False:
#    st.error("Username/password is incorrect")

#if authentication_status == None:
#    st.warning("Please enter your username and password")

#if authentication_status:
#    st.write(f"Welcome, {name}!")
#    q1, q2 = st.columns([7,1], vertical_alignment="center")
#    with q2:
        # Logout button
#        if st.button("Log Out", key="logout_button"):
#            authenticator.logout("Log Out")
#            authentication_status = None  # Reset auth state
#            st.rerun()  # Force redirect to login page
generation_config = {
    "candidate_count": 1,
    "max_output_tokens": 8192,
    "temperature": 0,
    "top_p": 0.95,
    "top_k": 1,
}
#api_key = st.secrets["genai"]["api_key"]
api_key = 'AIzaSyB5BOxYSGXekSac6H1ndOLJhPaMiP10qGE'
genai.configure(api_key = api_key)
safety_settings = [
    SafetySetting(
        category=SafetySetting.HarmCategory.HARM_CATEGORY_HATE_SPEECH,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
    ),
    SafetySetting(
        category=SafetySetting.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
    ),
    SafetySetting(
        category=SafetySetting.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
    ),
    SafetySetting(
        category=SafetySetting.HarmCategory.HARM_CATEGORY_HARASSMENT,
        threshold=SafetySetting.HarmBlockThreshold.BLOCK_MEDIUM_AND_ABOVE
    ),
]

def get_prompt(text, src, trg, tone, domain, instruction, mandatory_translations):
    if trg == 'bn':
        prompt = """
                    You are an expert Translator created by Lisan India. Your task is to translate texts **from {} to {}** with precision, ensuring correct writing direction (RTL or LTR). The text belongs to the **{}** domain and should be translated in a **{}** tone.

                    Your translation should be **accurate, natural, and professional**, following Bengali linguistic norms and industry best practices.

                    ---

                    ## **🔹 মূল অনুবাদ নির্দেশিকা (Key Translation Rules):**

                    ### **1. স্পষ্টতা ও প্রাকৃতিকতা (Clarity & Readability)**  
                    - অনুবাদ **স্বাভাবিক ও প্রাকৃতিক হওয়া উচিত**, যেন এটি বাংলায় মৌলিকভাবে লেখা হয়েছে।  
                    - **সহজ, স্পষ্ট ও ব্যবহারকারী-বান্ধব ভাষা** ব্যবহার করুন, যা বাংলা ভাষাভাষীদের কাছে সহজবোধ্য হবে।  
                    - **শব্দ-প্রতি-শব্দ অনুবাদ নয়**, বরং **অর্থ যথাযথভাবে প্রকাশ করুন**।  

                    ✅ **সঠিক:**  
                    - **English:** "Create beautiful designs from thousands of professional templates."  
                    - **Bengali:** "হাজারো পেশাদার টেমপ্লেট থেকে সুন্দর ডিজাইন তৈরি করুন।"  

                    ❌ **ভুল:**  
                    - "হাজারো পেশাদার টেমপ্লেট থেকে চমৎকার নকশা সৃষ্টি করুন।" (অপ্রাকৃতিক ও কঠিন শব্দ)  

                    ---

                    ### **2. বাক্য গঠন ও কণ্ঠস্বর (Sentence Structure & Voice)**  
                    - **সক্রিয় কণ্ঠ (Active Voice) ব্যবহার করুন**, যেখানে সম্ভব, প্যাসিভ এড়িয়ে চলুন।  
                    - **বাংলা ব্যাকরণের সঠিক ক্রম অনুসরণ করুন**, ইংরেজির মতো "Subject-Verb-Object" গঠন নয়।  

                    ✅ **সঠিক:**  
                    - **English:** "Here’s how to purchase your designs from Lisan."  
                    - **Bengali:** "এভাবেই আপনি Lisan থেকে আপনার ডিজাইন কিনতে পারেন।"  

                    ❌ **ভুল:**  
                    - "এই পদ্ধতি দেখানো হলো যার মাধ্যমে ডিজাইন কেনা যাবে।" (প্যাসিভ ও অপ্রাকৃতিক)  

                    ---

                    ### **3. ব্যবহারকারীকে সম্বোধন (User Addressing)**  
                    - **"আপনি" ব্যবহার করুন**, "তুমি" নয়, যাতে পেশাদার ও সম্মানজনক কণ্ঠ বজায় থাকে।  
                    - **সৌজন্যমূলক অথচ বন্ধুত্বপূর্ণ ভাষা ব্যবহার করুন**।  

                    ✅ **সঠিক:**  
                    - **English:** "Set up your account now."  
                    - **Bengali:** "আপনার অ্যাকাউন্ট এখন সেট করুন।"  

                    ❌ **ভুল:**  
                    - "তোমার অ্যাকাউন্ট এখন সেট করো।" (অতিরিক্ত অনানুষ্ঠানিক)  

                    ---

                    ### **4. প্রযুক্তিগত ও UI শর্তাবলী (Technical & UI Terms)**  
                    - **UI বোতামের নাম সর্বদা ক্রিয়ামূলক (verb form) হতে হবে।**  

                    ✅ **সঠিক:**  
                    - **English:** "Start" → **"শুরু করুন"**  
                    - **English:** "Continue" → **"চালিয়ে যান"**  

                    ❌ **ভুল:**  
                    - "এটি শুরু করুন" (অপ্রয়োজনীয় দীর্ঘ)  

                    - **প্রযুক্তিগত পরিভাষা (CPU, USB, PDF) বাংলায় অনুবাদ করবেন না।**  
                    - **উদাহরণ:** "Set the CPU host frequency" → **"CPU হোস্ট ফ্রিকোয়েন্সি সেট করুন"**  

                    ---

                    ### **5. বিরামচিহ্ন ও ফরম্যাটিং (Punctuation & Formatting)**  
                    - **বাংলা বিরামচিহ্ন অনুসরণ করুন** (যেমন, "এবং" বা "অথবা" এর আগে কমা নয়)।  
                    - **ব্র্যাকেট, প্রতীক (&, #, @) ও পথ নির্দেশিকা ইংরেজিতেই রাখুন।**  

                    ✅ **সঠিক:**  
                    - **English:** "Price (USD)"  
                    - **Bengali:** "মূল্য (USD)"  

                    ❌ **ভুল:**  
                    - "মূল্য (মার্কিন ডলার)" (USD অনুবাদ করা উচিত নয়)  

                    - **Boolean শব্দ ("AND", "OR", "IF") অনুবাদ করুন, তবে বড় হাতের অক্ষরে রাখুন।**  

                    ✅ **সঠিক:**  
                    - "IF" → "যদি"  
                    - "OR" → "অথবা"  
                    - "AND" → "এবং"  

                    ❌ **ভুল:**  
                    - "if" → "যদি" (এটি বড় হাতের অক্ষরে হওয়া উচিত)  

                    ---

                    ### **6. প্লেসহোল্ডার ও ভেরিয়েবলস (Handling Placeholders & Variables)**  
                    - **Curly brackets `{{}}` এর মধ্যে থাকা শব্দ অনুবাদ করবেন না।**  

                    ✅ **সঠিক:**  
                    - **English:** "Are you sure you want to delete {{row}}?"  
                    - **Bengali:** "আপনি কি নিশ্চিত যে আপনি {{row}} মুছতে চান?"  

                    ❌ **ভুল:**  
                    - "আপনি কি নিশ্চিত যে আপনি সারি {{row}} মুছতে চান?" ("row" অনুবাদ করা যাবে না)  

                    - **Nested placeholders `{{{{...{{{{...}}}}...}}}}` ক্ষেত্রে শুধুমাত্র অভ্যন্তরীণ অংশ অনুবাদ করুন।**  

                    ✅ **সঠিক:**  
                    - **English:** "The {{countEmails, plural, one {{email address already exists}} other {{email addresses already exist}}}} {{listOfEmails}}"  
                    - **Bengali:** "ডাটাবেসে {{countEmails, plural, one {{ইমেল ঠিকানা ইতিমধ্যে বিদ্যমান}} other {{ইমেল ঠিকানাগুলি ইতিমধ্যে বিদ্যমান}}}} {{listOfEmails}}"  

                    ---

                    ### **7. স্থানীয়করণ নিয়ম (Localization Rules)**  
                    - **তারিখ:** **MM/DD/YYYY → "তারিখ মাস বছর"**  
                    - **উদাহরণ:** "07/25/2016" → "২৫ জুলাই ২০১৬"  

                    - **সময় বিন্যাস:**  
                    - **১২-ঘণ্টা বিন্যাস (AM/PM) ব্যবহার করুন।**  
                    - **উদাহরণ:** "10:30 AM" → "১০:৩০ পূর্বাহ্ন"  

                    - **মুদ্রা রূপান্তর:**  
                    - **₹1,000 → "১,০০০ টাকা"**  
                    - **$1,000 → "১,০০০ ডলার"** (বিদেশি মুদ্রা অনুবাদ করবেন না)  

                    - **পরিমাপ একক:**  
                    - **"5 kg" → "৫ কেজি"**  
                    - **"50 cm" → "৫০ সেন্টিমিটার"**  

                    ---

                    ### **8. সাধারণ অনুবাদ পছন্দ (Common Translation Preferences)**  
                    - **বাংলা ভাষায় ইংরেজি গৃহীত শব্দ ব্যবহার করুন।**  
                    - **উদাহরণ:** "Follow" → **"ফলো করুন"**, না যে "অনুসরণ করুন"  
                    - **উদাহরণ:** "Check" → **"চেক করুন"**, না যে "পরীক্ষা করুন"  

                    - **অনাবশ্যক কঠিন বাংলা পরিভাষা এড়িয়ে চলুন।**  
                    - **উদাহরণ:** "Newspaper" → **"সংবাদপত্র"**, না যে "দৈনিক পত্রিকা"  

                    ---

                    ### **🔹 বাধ্যতামূলক অনুবাদ (Mandatory Translations):**  
                    {}
                    ### *Additional Instruction:*  
                    {}

                    ### **📌 অনুবাদ করার জন্য পাঠ্য (Text to Translate):**  
                    {}

                    ### **✅ আপনার অনুবাদ (Your Translation):** """.format(src, trg, domain, tone, mandatory_translations, instruction, text)

    elif trg == 'ar':
        prompt = """You are an expert Translator created by Lisan India. Your task is to translate texts *from {src} to {trg}* accurately with correct writing direction (RTL or LTR). 
                    The text belongs to the *{domain}* domain and should be translated in a *{tone}* tone.

                    ### *Important Instructions:*
                    1. *Strictly use the provided mandatory translations* if the first word is from {src} language and the other is in {trg} language.
                    2. *Do not modify* words that are replaced based on the dictionary.
                    3. *Ensure smooth, natural readability* while keeping accuracy.
                    4. *Keep action buttons as imperative verbs* (e.g., "Save", "Delete").
                    5. *Do not translate* anything inside curly brackets ({{}}). Example:  
                    - ❌ *Wrong:* "Consent accordé le {{date}}"  
                    - ✅ *Correct:* "Consent Granted on {{date}}"  
                    6. *Keep path identification elements in English*. Example:  
                    - ❌ *Wrong:* "sw2.stockage.appId doit être défini"  
                    - ✅ *Correct:* "sw2.storage.appId must be set"  
                    7. *Do not translate words without spaces*. Example:  
                    - "SW2 RecurrenceMonths" → should remain unchanged.  
                    8. *Translate Boolean indicators but keep them in CAPITAL LETTERS*. Example:  
                    - "IF", "OR", and "AND" should be translated but stay uppercase.  
                    9. *For nested curly brackets, translate only inner content*. Example:  
                    - "The {{countEmails, plural, one {{email address already exists}} other {{email addresses already exist}}} {{listOfEmails}}}"  
                    - *Only translate* "email address already exists" but *not* "countEmails" or "plural".

                    ### *Mandatory Translations (Do not modify these words):*  
                    {mandatory_translations}

                    ### *Instruction:*  
                    {instruction}

                    ### *Text to Translate:*  
                    {text}

                    ### *Your Translation:*  
                    """
    
    elif trg == 'hi':
        prompt = """You are an expert Translator created by Lisan India. Your task is to translate texts **from {} to {}** with precision, ensuring correct writing direction (RTL or LTR). The text belongs to the **{}** domain and should be translated in a **{}** tone.

                    Your translation should be **accurate, natural, and professional**, following Hindi linguistic norms and industry best practices.
                    ---

                    ## **Key Translation Rules:**

                    ### **1. Clarity & Readability**  
                    - Ensure the translation sounds **natural and locally adapted**, not like a literal translation from English.  
                    - Use **simple, clear, and user-friendly** language that feels natural to Hindi speakers.  
                    - **Avoid word-for-word translation**—instead, focus on conveying meaning.  

                    ✅ **Correct:**  
                    - **English:** "Create beautiful designs from thousands of professional templates."  
                    - **Hindi:** "हज़ारों पेशेवर टेम्पलेट से ख़ूबसूरत डिज़ाइन बनाएँ।"  

                    ❌ **Incorrect:**  
                    - "हज़ारों पेशेवर टेम्पलेट से सुंदर डिज़ाइन बनाइए।" (Too formal and unnatural)  

                    ---

                    ### **2. Sentence Structure & Voice**  
                    - **Prefer active voice over passive voice** for better clarity.  
                    - **Word order should follow Hindi grammar**, not English syntax.  

                    ✅ **Correct:**  
                    - **English:** "Here’s how to purchase your designs from Lisan."  
                    - **Hindi:** "आपकी डिज़ाइन को Lisan से खरीदने का तरीका यह है।"  

                    ❌ **Incorrect:**  
                    - "वह तरीका यहां प्रस्तुत है जिससे आपकी डिज़ाइन Lisan से खरीदी जा सकती है।" (Unnatural passive construction)  

                    ---

                    ### **3. User Addressing**  
                    - Always address the user as **"आप"** instead of **"तुम"** for a polite and professional tone.  
                    - Maintain **a conversational but professional tone** without being overly formal.  

                    ✅ **Correct:**  
                    - **English:** "Set up your account now."  
                    - **Hindi:** "अपना खाता अभी सेट करें।"  

                    ❌ **Incorrect:**  
                    - "तुम्हारा खाता अभी सेट करो।" (Too informal)  

                    ---

                    ### **4. Technical & UI Terms**  
                    - **Software button labels** should be translated using **verb forms** (command format).  

                    ✅ **Correct:**  
                    - **English:** "Start" → **"शुरू करें"**  
                    - **English:** "Continue" → **"आगे बढ़ें"**  

                    ❌ **Incorrect:**  
                    - "इसे शुरू कीजिए" (Too long and unnatural)  

                    - **Technical terms like "CPU", "USB", "PDF" should not be translated.**  
                    - **Example:** "Set the CPU host frequency" → **"CPU होस्ट फ़्रीक्वेंसी सेट करें"**  
                    - Do not attempt to translate "CPU" or "USB" into Hindi.  

                    ---

                    ### **5. Formatting & Punctuation**  
                    - **Follow Hindi punctuation rules** (do not place a comma before "और" or "या").  
                    - **Brackets, symbols (&, #, @), and path identifiers** should be kept unchanged.  

                    ✅ **Correct:**  
                    - **English:** "Price (USD)"  
                    - **Hindi:** "मूल्य (USD)"  

                    ❌ **Incorrect:**  
                    - "मूल्य (अमेरिकी डॉलर)" (Adding unnecessary translation for "USD")  

                    - **Boolean terms ("AND", "OR", "IF") should be translated but kept in CAPITAL LETTERS.**  

                    ✅ **Correct:**  
                    - "IF" → "यदि"  
                    - "OR" → "या"  
                    - "AND" → "और"  

                    ❌ **Incorrect:**  
                    - "if" → "यदि" (Should be capitalized)  

                    ---

                    ### **6. Handling Placeholders & Variables**  
                    - **Do not translate text inside curly brackets `{{}}`.**  

                    ✅ **Correct:**  
                    - **English:** "Are you sure you want to delete {{row}}?"  
                    - **Hindi:** "क्या आप वाकई {{row}} को हटाना चाहते हैं?"  

                    ❌ **Incorrect:**  
                    - "क्या आप वाकई पंक्ति {{row}} को हटाना चाहते हैं?" (Translating "row" incorrectly)  

                    - **For nested curly brackets `{{...{{...}}...}}`**, translate only the **inner content**, not the placeholders.  

                    ✅ **Correct:**  
                    - **English:** "The {{countEmails, plural, one {{email address already exists}} other {{email addresses already exist}}}} {{listOfEmails}}"  
                    - **Hindi:** "डेटाबेस में {{countEmails, plural, one {{ईमेल पता पहले से मौजूद है}} other {{ईमेल पते पहले से मौजूद हैं}}}} {{listOfEmails}}"  

                    ❌ **Incorrect:**  
                    - "डेटाबेस में {{गिनतीईमेल, बहुवचन, एक {{ईमेल पता पहले से मौजूद है}} अन्य {{ईमेल पते पहले से मौजूद हैं}}}} {{सूचीईमेल}}" (Translating outer placeholders)  

                    ---

                    ### **7. Localization Rules**  
                    - **Dates:** Convert **MM/DD/YYYY → Date Month Year.**   
                    - **Example:** "07/25/2016 – 25 जुलाई 2016." (strictly) 

                    - **Time Format:** Use **12-hour format** unless for railway/airline schedules.  
                    - **Example:** "10:30 AM" → "10:30 पूर्वाह्न"  

                    - **Currency Handling:**  
                    - **₹1,000 → "1,000 रु."** (Keep Indian Rupees format)  
                    - **$1,000 → "1,000 डॉलर"** (Do not convert foreign currencies)  

                    - **Measurement Units:** Use standard Hindi terms.  
                    - **Example:** "5 kg" → **"5 किलोग्राम"**  
                    - **Example:** "50 cm" → **"50 सेंटीमीटर"**  

                    ---

                    ### **8. Common Translation Preferences**  
                    - **Use borrowed English words** where appropriate.  
                    - **Example:** "Follow" → **"फ़ॉलो करें"**, not "अनुगमन करें"  
                    - **Example:** "Check" → **"चेक करें"**, not "जाँच करें"  

                    - **Avoid Sanskritized Hindi** if a simpler Urdu/Persian term is more common.  
                    - **Example:** "Newspaper" → **"अख़बार"**, not "समाचार-पत्र"  

                    - **Avoid literal translations.**  
                    - **Example:** "Scientific and award-winning" → **"पुरस्कृत और विज्ञान पर आधारित"**, not "वैज्ञानिक एवं पुरस्कार विजेता"  

                    ---

                    ### **Mandatory Translations (Do not modify these words):**  
                    {}
                    ### **Additional Instruction:**  
                    {}
                    ### **Text to Translate:**  
                    {}

                    ### **Your Translation:**  
                    """.format(src, trg, domain, tone, mandatory_translations, instruction, text)
    else:
        prompt = """You are an expert Translator create by Lisan India. Your task is to translate texts **from {src} to {trg}** accurately with correct writing diraection (RTL or LTR). 
                The text belongs to the **{domain}** domain and should be translated in a **{tone}** tone.

                ### **Important Instructions:**
                1. **Strictly use the provided mandatory translations** if 1st word is from {src} language and other is in {trg} language.
                2. **Do not modify** words that are replaced based on the dictionary.
                3. **Ensure smooth, natural readability** while keeping accuracy.

                ### **Mandatory Translations (Do not modify these words):**  
                {mandatory_translations}

                ### **Instruction:**  
                {instruction}

                ### **Text to Translate:**  
                {text}

                ### **Your Translation:**  
                """
    return prompt
        

def generate(text, src, trg, llm_model, tone='formal', domain='Healthcare', instruction='0'):
    # Initialize Vertex AI with project and location from secrets
    service_account_info = st.secrets["gcp_service_account"]
    credentials_path = os.path.abspath("service_account_key.json")
    vertexai.init(
        project = service_account_info["project_id"],
        location = "us-central1",
        credentials = credentials_path,
    )
    
    model = GenerativeModel(
        model_name=llm_model
    )

    # Generate content
    responses = model.generate_content(
        [f'You are an expert Translator. You are tasked to translate documents from {src} to {trg}. \
        Please provide an accurate translation of this text which is from {domain} and return translation text only, considering the {tone} \
        Instruction: {instruction} \
        :{text}'],
        generation_config=generation_config,
        safety_settings=safety_settings,
    )

    return responses.candidates[0].content.parts[0].text

def convert_df_to_excel(df):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False)
    return buffer.getvalue()  # return raw binary for download

def translate_text(text, src, trg, llm_model, tone, domain, instruction, mandatory_translations = 'None'):

    # Stronger Prompt Template
    prompt = '''You are an expert Translator create by Lisan India. Your task is to translate texts **from {} to {}** accurately with correct writing diraection (RTL or LTR). 
                The text belongs to the **{}** domain and should be translated in a **{}** tone.

                ### **Important Instructions:**
                1. **Strictly use the provided mandatory translations** if 1st word is from {} language and other is in {} language.
                2. **Do not modify** words that are replaced based on the dictionary.
                3. **Ensure smooth, natural readability** while keeping accuracy.

                ### **Mandatory Translations (Do not modify these words):**  
                {}

                ### **Instruction:**  
                {}

                ### **Text to Translate:**  
                {}

                ### **Your Translation:**  
                '''.format(src, trg, domain, tone, src, trg, mandatory_translations, instruction, text)
    model = genai.GenerativeModel(llm_model)
    response = model.generate_content(prompt)
    
    return response.text.strip()

def mqxliff_to_df(file_path, source_col= 'Source', target_col = 'Target'):
    tree = etree.parse(file_path)
    root = tree.getroot()
    nsmap = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}
    rows = []

    for file_element in root.findall('xliff:file', namespaces=nsmap):
        for body in file_element.findall('xliff:body', namespaces=nsmap):
            for trans_unit in body.findall('xliff:trans-unit', namespaces=nsmap):
                unit_id = trans_unit.get('id')
                source = trans_unit.find('xliff:source', namespaces=nsmap)
                target = trans_unit.find('xliff:target', namespaces=nsmap)
                source_text = source.text if source is not None else ''
                target_text = target.text if target is not None else ''
                rows.append({
                    'ID': unit_id,
                    source_col: source_text,
                    target_col: target_text
                })

    # Convert to DataFrame
    df = pd.DataFrame(rows)

    # Save to Excel
    #df.to_excel(output_path, index=False)
    #print(f"✅ Excel saved to: {output_path}")
    return df

def sdlxliff_to_df(file_path, source_col='Source', target_col='Target', debug=False):
    tree = etree.parse(file_path)
    root = tree.getroot()

    nsmap = {
        'xliff': 'urn:oasis:names:tc:xliff:document:1.2',
        'sdl': 'http://sdl.com/FileTypes/SdlXliff/1.0'
    }

    rows = []
    trans_units = root.findall('.//xliff:trans-unit', namespaces=nsmap)
    if debug:
        print("Found trans-units:", len(trans_units))

    for unit in trans_units:
        unit_id = unit.get("id")

        # Try <mrk> inside <seg-source>
        source_mrks = unit.findall('.//xliff:seg-source//xliff:mrk', namespaces=nsmap)
        if source_mrks:
            source_text = ' '.join([''.join(m.itertext()).strip() for m in source_mrks])
        else:
            # Fallback to <source>
            source_el = unit.find('.//xliff:source', namespaces=nsmap)
            source_text = ''.join(source_el.itertext()).strip() if source_el is not None else ''

        # Try <mrk> inside <target>
        target_mrks = unit.findall('.//xliff:target//xliff:mrk', namespaces=nsmap)
        if target_mrks:
            target_text = ' '.join([''.join(m.itertext()).strip() for m in target_mrks])
        else:
            # Fallback to <target>
            target_el = unit.find('.//xliff:target', namespaces=nsmap)
            target_text = ''.join(target_el.itertext()).strip() if target_el is not None else ''

        # Include rows that have either
        if source_text or target_text:
            rows.append({
                'ID': unit_id,
                source_col: source_text,
                target_col: target_text
            })
        elif debug:
            print(f"❌ Skipping unit {unit_id} — no text found")
            print("Raw <source>:", etree.tostring(unit.find('.//xliff:source', namespaces=nsmap), encoding='unicode'))
            print("Raw <target>:", etree.tostring(unit.find('.//xliff:target', namespaces=nsmap), encoding='unicode'))
            print("-" * 50)


    df = pd.DataFrame(rows)
    if debug and df.empty:
        print("⚠️ No usable data found.")
    return df


def clean_xml_text(text):
    if not isinstance(text, str):
        return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

def df_to_xliff(df, xliff_input_path, output_path):
    tree = etree.parse(xliff_input_path)
    root = tree.getroot()
    nsmap = {'xliff': 'urn:oasis:names:tc:xliff:document:1.2'}

    id_to_target = dict(zip(df["ID"], df["Target"]))

    for trans_unit in root.findall('.//xliff:trans-unit', namespaces=nsmap):
        unit_id = trans_unit.get('id')
        if unit_id in id_to_target:
            # Find or create <target>
            target_el = trans_unit.find('xliff:target', namespaces=nsmap)
            if target_el is None:
                target_el = etree.SubElement(trans_unit, 'target', attrib={"state": "translated"})
            else:
                target_el.attrib["state"] = "translated"

            target_el.text = clean_xml_text(id_to_target[unit_id])

    tree.write(output_path, encoding='utf-8', pretty_print=True, xml_declaration=True)
    #print(f"✅ Clean .xliff with updated targets saved to: {output_path}")
    return output_path
def get_transcript(audio_file, audio_language='unknown'): 
    url = "https://api.sarvam.ai/speech-to-text"

    files = {
        "file": ('audio.wav', open(audio_file,'rb'), "audio/wav")  # Convert to WAV format
    }

    data = {
        "language_code": audio_language,
        "model": "saarika:v2",
        "with_diarization": "true",
        "with_timestamps": "true"
    }

    headers = {
        "api-subscription-key": "5a73b765-cbce-43bd-8080-c7430ce4d961"  # Replace with your API key
    }

    response = requests.post(url, files=files, data=data, headers=headers)

    return response


def split_audio(audio_file, segment_length=149*60*100, output_dir="audio_segments"):
    os.makedirs(output_dir, exist_ok=True)  # Create directory if not exists

    audio = AudioSegment.from_file(audio_file)
    total_duration = len(audio)  # Duration in milliseconds

    if total_duration <= segment_length:
        segment_path = os.path.join(output_dir, f"segment.wav")
    else:
        segments = []
        for i, start in enumerate(range(0, total_duration, segment_length)):
            end = min(start + segment_length, total_duration)
            segment = audio[start:end]
            
            segment_path = os.path.join(output_dir, f"segment_{i}.wav")
            segment.export(segment_path, format="wav")  # Save segment
            segments.append(segment_path)

    return output_dir

def generate_NMT(strs_to_translate: List[str], src: str, tgt: str
) -> translate.TranslationServiceClient:
    """Translating Text."""

    client = translate.TranslationServiceClient()

    location = "us-central1"

    parent = f"projects/lisanai/locations/{location}"

    # Translate text from en to fr
    response = client.translate_text(
        request={
            "parent": parent,
            "contents": strs_to_translate,
            "mime_type": "text/plain",  # mime types: text/plain, text/html
            "source_language_code": src,
            "target_language_code":  tgt,
        }
    )

    return [text.translated_text for text in response.translations]  

def load_lottiefile(filepath: str):
        with open(filepath, "r") as f:
            return json.load(f)


def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

def df_to_json(df, source_col = 'Source', target_col= 'Target'):
    return df[[source_col]].rename(columns={source_col: 'text'}).to_dict(orient='records')

#def df_to_json(df, source_col="source"):
#    if source_col not in df.columns:
#        raise ValueError(f"Expected column '{source_col}' not found in DataFrame. Found columns: {df.columns.tolist()}")
#    return df[[source_col]].rename(columns={source_col: 'text'}).to_dict(orient='records')
def df_to_json_single_quotes(df, source_col='Source', target_col='Target'):
    # Convert to JSON string with proper escaping
    data = df[[source_col]].rename(columns={source_col: 'text'}).to_dict(orient='records')
    
    # Convert to string using repr to get single quotes
    return repr(data)
    '''json_str = json.dumps(
        df[[source_col]].rename(columns={source_col: 'text'}).to_dict(orient='records')
    )
    # Replace outer double quotes with single quotes, but keep inner ones
    # Trick: use regex to only replace outer ones
    import re
    fixed = re.sub(r'"([^"]*?)"\s*:', r"'\1':", json_str)  # keys
    fixed = re.sub(r':\s*"((?:[^\"\\]|\\.)*?)"', r": '\1'", fixed)  # string values
    return fixed'''


def translate_json(text_json, target, source, tone, domain, instruction, mandatory_translations = 'None', source_col = 'Source', target_col= 'Target'):
    llm_model_config = {'temperature': 0.1, 'top_p': 1, 'top_k': 40,
                        'max_output_tokens': 100000000, 'response_mime_type': 'application/json'}
    llm_model = genai.GenerativeModel(model_name = 'gemini-2.0-flash', generation_config = llm_model_config)

    res_schema = '''[{{
                '{}': 'string, required — the same as in the input',
                '{}': 'string, required — the translated version of the "text" field',
                '{}': 'string, required — the same as in the input',
                '{}': 'empty string, if the text is null or empty',
                .
                .
                }}]'''.format(source_col, target_col, source_col, target_col)
    prompt = '''You are an expert Translator created by Lisan India. Your task is to translate texts **from {} to {}** accurately with correct writing diraection (RTL or LTR).
                The text belongs to the **{}** domain and should be translated in a **{}** tone.

                ### **Important Instructions:**
                1. **Strictly use the provided mandatory translations** if 1st word is from {} language and other is in {} language.
                2. **Do not modify** words that are replaced based on the dictionary.
                3. **Ensure smooth, natural readability** while keeping accuracy.


                ### **Mandatory Translations (Do not modify these words):**
                {}

                ### **Instruction:**
                {}

                ##Response Schema - the response should be strictly in this format
                {}
                ### **Text to Translate:**
                {}
                ### **Your Translation:**
                '''.format(source, target, domain, tone, source, target, mandatory_translations, instruction, res_schema, text_json)

    response = llm_model.generate_content(prompt)
    raw = response.text
    # Ensure the raw response is treated as a string before regex
    json_str = eval(raw)


    return json_str # Return as a list of dictionaries

def batch_translate_json(json_payload, source, target, tone, domain, instruction, mandatory_translations,  batch_size=50, delay=2):
    result = []
    for i in range(0, len(json_payload), batch_size):
        batch = json_payload[i:i+batch_size]
        translated_batch = translate_json(batch, target, source, tone, domain, instruction, mandatory_translations)
        result.extend(translated_batch)
        time.sleep(delay)  # Throttle to avoid rate limits
    return result

def batch_translate_df(df, source, target, tone, domain, instruction, mandatory_translations, source_col, target_col, batch_size=50, delay=2):
    result_df = pd.DataFrame()
    for i in range(0, len(df), batch_size):
        batch = df.iloc[i:i+batch_size]
        json_payload = df_to_json_single_quotes(batch, source_col, target_col)
        translated_json = translate_json(json_payload, target, source, tone, domain, instruction, mandatory_translations, source_col, target_col)
        batch_result = pd.DataFrame(translated_json)
        result_df = pd.concat([result_df, batch_result], ignore_index=True)
        
        time.sleep(delay)  # Throttle requests to avoid rate limits

    return result_df

def inject_translations_to_xliff(uploaded_file, po_path='result.po'):
    # If it's an UploadedFile, save it to temp
    if hasattr(uploaded_file, "read"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
            tmp.write(uploaded_file.read())
            original_path = tmp.name
    else:
        # Assume it's already a filesystem path
        original_path = uploaded_file

    base, ext = os.path.splitext(original_path)
    output_path = f"{base}_result{ext}"

    # Load translations
    po = polib.pofile(po_path)
    translations = {
        entry.msgid.strip(): entry.msgstr.strip()
        for entry in po if entry.msgstr.strip()
    }

    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(original_path, parser)
    root = tree.getroot()

    nsmap = {}
    if None in root.nsmap:
        nsmap['ns'] = root.nsmap[None]
    for k, v in root.nsmap.items():
        if k:
            nsmap[k] = v

    trans_units = root.xpath('//ns:trans-unit' if 'ns' in nsmap else '//trans-unit', namespaces=nsmap)

    for tu in trans_units:
        source_elem = tu.find('.//ns:source' if 'ns' in nsmap else './/source', namespaces=nsmap)
        target_elem = tu.find('.//ns:target' if 'ns' in nsmap else './/target', namespaces=nsmap)

        if source_elem is not None:
            source_text = (source_elem.text or "").strip()
            if source_text in translations:
                if target_elem is None:
                    target_tag = '{%s}target' % nsmap.get('ns', '') if 'ns' in nsmap else 'target'
                    target_elem = etree.SubElement(tu, target_tag)
                target_elem.text = translations[source_text]

    tree.write(output_path, pretty_print=True, xml_declaration=True, encoding="utf-8")
    print(f"[OK] Injected translations into: {output_path}")
    return output_path
def json_to_po(translated_list, original_po_path = 'cleaned.po', output_po_path = 'result.po'):
    # Convert your list of dicts to a lookup dict for fast access
    translation_map = {item["Source"]: item["Target"] for item in translated_list}

    po = polib.pofile(original_po_path)

    for entry in po:
        src = entry.msgid.strip()
        if src in translation_map:
            entry.msgstr = translation_map[src] or ""  # assign translated text

    po.save(output_po_path)
    print(f"[DONE] Saved translated PO file: {output_po_path}")

def po_to_json_payload(po_path = 'cleaned.po'):
    po = polib.pofile(po_path)

    # Build list of dicts
    data = []
    for entry in po:
        if entry.msgid.strip():  # skip empty
            data.append({"text": entry.msgid.strip()})

    # Return JSON string like your Excel function
    return repr(data)

def convert_to_po(po_path='cleaned.po', xliff_path = 'cleaned.xliff'):
    # Convert to PO
    result = subprocess.run(
        ["xliff2po", xliff_path, po_path],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )
    print(result.stdout)
    if result.returncode != 0:
        print("[ERROR] Conversion failed:\n", result.stderr)
    else:
        print(f"[SUCCESS] Converted to: {po_path}")


def clean_custom_xliff(input_path, output_path = 'cleaned.xliff'):
    """
    Cleans vendor-specific XLIFF formats (SDLXLIFF, MQXLIFF, etc.) into plain XLIFF.
    Removes namespaces, vendor-specific attributes, and extra metadata.
    """
    parser = etree.XMLParser(remove_blank_text=True)
    tree = etree.parse(input_path, parser)
    root = tree.getroot()

    # Remove namespaces
    for elem in root.iter():
        if not hasattr(elem.tag, 'find'):
            continue
        i = elem.tag.find('}')
        if i >= 0:
            elem.tag = elem.tag[i+1:]  # Strip namespace
        # Remove attributes that aren't standard XLIFF attributes
        elem.attrib.clear()

    etree.cleanup_namespaces(root)
    # Save clean file
    tree.write(output_path, pretty_print=True, xml_declaration=True, encoding="utf-8")
    print(f"[CLEANED] Saved plain XLIFF: {output_path}")



def translate_excel(input_file, source_col_name, target_col_name):
    df = pd.read_excel(input_file)
    json_payload = df_to_json(df)
    translated_json = translate_json(json_payload, source_col_name, target_col_name)
    df_res = pd.DataFrame(translated_json)
    df_res.rename(columns={'translated': target_col_name}, inplace=True)
    df_res.to_excel('result_file.xlsx', index=False)
#print("Fiile Saved to " + result_file)
c1, c2, c3 = st.columns([2,5,1], vertical_alignment="center")
lottie_hello = load_lottieurl("https://lottie.host/057e0efe-27c7-4397-840c-f1f25b8a682a/6Dw9TLkyW5.json")
with c2:
    st_lottie(
        lottie_hello,
        speed=1,
        reverse=False,
        loop=True,
        quality="low", # medium ; high
        #renderer="svg", # canvas
        height=300,
        width=300,
        key=None,

    )
a1, a2, a3 = st.columns([1,3,1], vertical_alignment="center")
with a2:
    #st.title("Your AI for your Documents")
    #st.markdown("<h1 style='text-align: center;'>Hello{name}</h1>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align: center;'>Lisan AI</h1>", unsafe_allow_html=True)

audio_on = st.toggle("Audio")

if audio_on:
    audio_file = st.file_uploader("Upload an audio file", type=["mp3", "wav", "wave", "x-wav", "mpeg"])

    audio_language_dict = {
        "Unknown": "unknown",
        "Hindi": "hi-IN",
        "Bengali": "bn-IN",
        "Kannada": "kn-IN",
        "Malayalam": "ml-IN",
        "Marathi": "mr-IN",
        "Odia": "od-IN",
        "Punjabi": "pa-IN",
        "Tamil": "ta-IN",
        "Telugu": "te-IN",
        "English (India)": "en-IN",
        "Gujarati": "gu-IN"
    }

    language_opts = tuple(audio_language_dict.keys())
    audio_language =st.selectbox(
            "Select the input Audio Language",
            language_opts,
            index=0
        )
    
    if st.button("Get Transcript"):                    
        if audio_file is not None:
            st.audio(audio_file, format="audio/wav")

            # Convert Streamlit file to BytesIO
            audio_bytes = BytesIO(audio_file.read())

            # Split the audio into segments
            output_dir = split_audio(audio_bytes)
            responses = []
            start_time, end_time, speaker_id, transcript =[], [], [], []
            for audio in output_dir:
                response = get_transcript(audio, audio_language_dict[audio_language])
                
                for i in eval(response.text)['diarized_transcript']['entries']:
                        start_time.append(i['start_time_seconds'])
                        end_time.append(i['end_time_seconds'])
                        speaker_id.append(i['speaker_id'])
                        transcript.append(i['transcript'])
            data = {
                    "Start Time": start_time,
                    "End Time": end_time,
                    "Speaker IDs": speaker_id,
                    "Transcripts": transcript
                }
            df = pd.DataFrame(data)
            df.to_excel('Transcript.xlsx', index=True)
            with open('Transcript.xlsx', "rb") as template_file:
                template_byte = template_file.read()

            st.download_button(label="Download Transcript",
                                data=template_byte,
                                file_name="Transcript.xlsx",
                                mime='application/octet-stream')
    
else:
    b1, b2 = st.columns([1,1], vertical_alignment="center")
    languages = {
        "Abkhaz": "ab",
        "Acehnese": "ace",
        "Acholi": "ach",
        "Afrikaans": "af",
        "Albanian": "sq",
        "Alur": "alz",
        "Amharic": "am",
        "Arabic": "ar",
        "Armenian": "hy",
        "Assamese": "as",
        "Awadhi": "awa",
        "Aymara": "ay",
        "Azerbaijani": "az",
        "Balinese": "ban",
        "Bambara": "bm",
        "Bashkir": "ba",
        "Basque": "eu",
        "Batak Karo": "btx",
        "Batak Simalungun": "bts",
        "Batak Toba": "bbc",
        "Belarusian": "be",
        "Bemba": "bem",
        "Bengali": "bn",
        "Betawi": "bew",
        "Bhojpuri": "bho",
        "Bikol": "bik",
        "Bosnian": "bs",
        "Breton": "br",
        "Bulgarian": "bg",
        "Buryat": "bua",
        "Cantonese": "yue",
        "Catalan": "ca",
        "Cebuano": "ceb",
        "Chichewa (Nyanja)": "ny",
        "Chinese (Simplified)": "zh-CN or zh (BCP-47)",
        "Chinese (Traditional)": "zh-TW (BCP-47)",
        "Chuvash": "cv",
        "Corsican": "co",
        "Crimean Tatar": "crh",
        "Croatian": "hr",
        "Czech": "cs",
        "Dari": "Dari",
        "Danish": "da",
        "Dinka": "din",
        "Divehi": "dv",
        "Dogri": "doi",
        "Dombe": "dov",
        "Dutch": "nl",
        "Dzongkha": "dz",
        "English": "en",
        "Esperanto": "eo",
        "Estonian": "et",
        "Ewe": "ee",
        "Fijian": "fj",
        "Filipino (Tagalog)": "fil or tl",
        "Finnish": "fi",
        "French": "fr",
        "French (French)": "fr-FR",
        "French (Canadian)": "fr-CA",
        "Frisian": "fy",
        "Fulfulde": "ff",
        "Ga": "gaa",
        "Galician": "gl",
        "Ganda (Luganda)": "lg",
        "Georgian": "ka",
        "German": "de",
        "Greek": "el",
        "Guarani": "gn",
        "Gujarati": "gu",
        "Haitian Creole": "ht",
        "Hakha Chin": "cnh",
        "Hausa": "ha",
        "Hawaiian": "haw",
        "Hebrew": "iw or he",
        "Hiligaynon": "hil",
        "Hindi": "hi",
        "Hmong": "hmn",
        "Hungarian": "hu",
        "Hunsrik": "hrx",
        "Icelandic": "is",
        "Igbo": "ig",
        "Iloko": "ilo",
        "Indonesian": "id",
        "Irish": "ga",
        "Italian": "it",
        "Japanese": "ja",
        "Javanese": "jw or jv",
        "Kannada": "kn",
        "Kapampangan": "pam",
        "Kazakh": "kk",
        "Khmer": "km",
        "Kiga": "cgg",
        "Kinyarwanda": "rw",
        "Kituba": "ktu",
        "Konkani": "gom",
        "Korean": "ko",
        "Krio": "kri",
        "Kurdish (Kurmanji)": "ku",
        "Kurdish (Sorani)": "ckb",
        "Kyrgyz": "ky",
        "Lao": "lo",
        "Latgalian": "ltg",
        "Latin": "la",
        "Latvian": "lv",
        "Ligurian": "lij",
        "Limburgan": "li",
        "Lingala": "ln",
        "Lithuanian": "lt",
        "Lombard": "lmo",
        "Luo": "luo",
        "Luxembourgish": "lb",
        "Macedonian": "mk",
        "Maithili": "mai",
        "Makassar": "mak",
        "Malagasy": "mg",
        "Malay": "ms",
        "Malay (Jawi)": "ms-Arab",
        "Malayalam": "ml",
        "Maltese": "mt",
        "Maori": "mi",
        "Marathi": "mr",
        "Meadow Mari": "chm",
        "Meiteilon (Manipuri)": "mni-Mtei",
        "Minang": "min",
        "Mizo": "lus",
        "Mongolian": "mn",
        "Myanmar (Burmese)": "my",
        "Ndebele (South)": "nr",
        "Nepalbhasa (Newari)": "new",
        "Nepali": "ne",
        "Northern Sotho (Sepedi)": "nso",
        "Norwegian": "no",
        "Nuer": "nus",
        "Occitan": "oc",
        "Odia (Oriya)": "or",
        "Oromo": "om",
        "Pangasinan": "pag",
        "Papiamento": "pap",
        "Pashto": "ps",
        "Persian": "fa",
        "Polish": "pl",
        "Portuguese": "pt",
        "Portuguese (Portugal)": "pt-PT",
        "Portuguese (Brazil)": "pt-BR",
        "Punjabi": "pa",
        "Punjabi (Shahmukhi)": "pa-Arab",
        "Quechua": "qu",
        "Romani": "rom",
        "Romanian": "ro",
        "Rundi": "rn",
        "Russian": "ru",
        "Samoan": "sm",
        "Sango": "sg",
        "Sanskrit": "sa",
        "Scots Gaelic": "gd",
        "Serbian": "sr",
        "Sesotho": "st",
        "Seychellois Creole": "crs",
        "Shan": "shn",
        "Shona": "sn",
        "Sicilian": "scn",
        "Silesian": "szl",
        "Sindhi": "sd",
        "Sinhala (Sinhalese)": "si",
        "Slovak": "sk",
        "Slovenian": "sl",
        "Somali": "so",
        "Spanish": "es",
        "Sundanese": "su",
        "Swahili": "sw",
        "Swati": "ss",
        "Swedish": "sv",
        "Tajik": "tg",
        "Tamil": "ta",
        "Tatar": "tt",
        "Telugu": "te",
        "Tetum": "tet",
        "Thai": "th",
        "Tigrinya": "ti",
        "Tsonga": "ts",
        "Tswana": "tn",
        "Turkish": "tr",
        "Turkmen": "tk",
        "Twi (Akan)": "ak",
        "Ukrainian": "uk",
        "Urdu": "ur",
        "Uyghur": "ug",
        "Uzbek": "uz",
        "Vietnamese": "vi",
        "Welsh": "cy",
        "Xhosa": "xh",
        "Yiddish": "yi",
        "Yoruba": "yo",
        "Yucatec Maya": "yua",
        "Zulu": "zu"
    }
    keys_lang = tuple(languages.keys())


    llm_model =st.selectbox(
            "Model Selection",
            ("gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-flash-002", "gemini-1.5-pro-002"),
            index=0
        )
    # Layout with columns
    # "gemini-1.5-flash-001", "gemini-1.5-pro-001", "gemini-1.0-pro-001"
    b1, b2, b3 = st.columns([1, 0.5, 1])

    with b1:
        source = st.selectbox(
            "Source Language",
            keys_lang,
            index=48,
            key="source_lang"
        )
    with b3:
        target = st.selectbox(
            "Target Language",
            keys_lang,
            index=7,
            key="target_lang"
        )
    x1, x2= st.columns([1, 1])

    with x1:
        tone = st.selectbox(
            "Translation Tone",
            ('Normal', 'Formal', 'Infomal'),
            index=0,
            key="lang_tone"
        )
    with x2:
        domain = st.selectbox(
            "Translation Domain",
            ('General', 'Financial', 'Educational', 'Healthcare', 'Technology', 'Business'),
            index=0,
            key="lang_domain"
        )
    
    instruction = st.text_area("Translation Instruction", 'None', height=250)

    uploaded_file = st.file_uploader("Upload approved list of terms (Bilingual)", type=["xlsx"])

    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file, header=None)  # Read without headers

        # Display first few rows for reference
        #st.write("Preview of Uploaded File:", df.head())

        if len(df.columns) >= 2:
            # Let user select columns based on position
            word_col = 0 #st.number_input("Select column index for 'Words' (starting from 0)", min_value=0, max_value=len(df.columns)-1, value=0, step=1)
            translation_col = 1 #st.number_input("Select column index for 'Translations' (starting from 0)", min_value=0, max_value=len(df.columns)-1, value=1, step=1)

            # Extract selected columns
            df_selected = df.iloc[:, [word_col, translation_col]]
            df_selected.columns = ["word", "translation"]  # Rename for consistency
            
            st.write("Preview of Uploaded File:", df_selected.head())
            mandatory_translations = "\n".join([f"- {row['word']} → {row['translation']}" for _, row in df_selected.iterrows()])
    
        else:
            mandatory_translations = 'None'

    else:
        mandatory_translations = 'None'
    
    on = st.toggle("Text File")

    if on:
        uploaded_file = st.file_uploader("Upload a file", type=["txt", "xlsx","csv", "pdf", "docx", "xliff", "mqxliff", "sdlxliff", "mxliff"])
        source_col = st.text_area("Source Column Name", 'Source')
        target_col = st.text_area("Target Column Name", 'Target')

        if uploaded_file is not None:
            filename = uploaded_file.name
            file_extension = filename.split(".")[-1]

            if file_extension in ["txt", "csv"]:
                # Read text or CSV file
                stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
                text = stringio.read()
                st.write("Extracted Text:", text)
            
            elif file_extension == "xliff" or file_extension == "mqxliff" or file_extension == "sdlxliff" or file_extension == "mxliff"  :
                input_path = uploaded_file
                clean_custom_xliff(input_path)
                convert_to_po()
                json_payload = po_to_json_payload()
                text_json = eval(json_payload)


            elif file_extension == "xlsx":
                df_excel = pd.read_excel(uploaded_file)  # Load all sheets
                st.dataframe(df_excel)
            elif file_extension == "pdf":
                # Read PDF file
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
                st.write("Extracted Text:", text)

            elif file_extension == "docx":
                # Read DOCX file
                doc = docx.Document(uploaded_file)
                extracted_data = []
                for element in doc.element.body:
                    if isinstance(element, CT_P):  # If it's a paragraph
                        para = element.xpath(".//w:t")  # Extract text
                        if para:
                            extracted_data.append(" ".join([t.text for t in para if t.text.strip()]))  # Join text
                    
                    elif isinstance(element, CT_Tbl):  # If it's a table
                        table_data = []
                        table = next(tbl for tbl in doc.tables if tbl._element is element)  # Find corresponding table
                        for row in table.rows:
                            table_data.append("\t".join([cell.text.strip() for cell in row.cells]))  # Join table row with tabs
                        extracted_data.append("\n".join(table_data))  # Join all rows with newlines
                
                text =  "\n\n".join(extracted_data) 
                st.write("Extracted Text:", text)

            else:
                st.error("Unsupported file format")

            if st.button("Translate"):
                if file_extension in ["txt", "csv"]:
                    # Save as text file
                    st.download_button("Download Translated File", data=translated_text, file_name=f"Translated_{filename}")
                elif file_extension == "xlsx":
                    if len(df_excel) <= 75:
                        json_payload = df_to_json_single_quotes(df_excel, source_col, target_col)
                        translated_json = translate_json(json_payload, target, source, tone, domain, instruction, mandatory_translations, source_col, target_col)
                        df_res = pd.DataFrame(translated_json)
                    else:
                        df_res = batch_translate_df(df_excel, source, target, tone, domain, instruction, mandatory_translations, source_col, target_col)
                    #df_res.rename(columns={'translated': target_col}, inplace=True)
                    st.dataframe(df_res)
                    
                    excel_data = convert_df_to_excel(df_res)
                    
                    # Create download button
                    st.download_button(
                        label="📥 Download as Excel",
                        data=excel_data,
                        file_name="Result.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                elif file_extension == "docx":
                    # Save as DOCX file
                    translated_doc = docx.Document()
                    translated_text = translate_text(text, languages[source], languages[target], llm_model, tone, domain, instruction, mandatory_translations)
                    st.write(translated_text)
                    translated_doc.add_paragraph(translated_text)
                    docx_buffer = BytesIO()
                    translated_doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.download_button("Download Translated File", data=docx_buffer, file_name=f"Translated_{filename}")

                elif file_extension == "xliff" or file_extension == "mqxliff" or file_extension == "sdlxliff" or file_extension == "mxliff":
                    if len(text_json)> 50:
                        result = batch_translate_json(text_json, source, target, tone, domain, instruction, mandatory_translations)
                    else:
                        result = translate_json(text_json, target, source, tone, domain, instruction, mandatory_translations)
                    json_to_po(result)
                    output_path = inject_translations_to_xliff(input_path)
                    
                    with open(output_path, "rb") as f:
                        xliff_data = f.read()

                    # Streamlit download button
                    st.download_button(
                        label="📥 Download Translated XLIFF",
                        data=xliff_data,
                        file_name=output_path,
                        mime="application/xml"
                    )
                elif file_extension == "pdf":
                    # **Exception:** Save as TXT instead of PDF
                    st.download_button("Download Translated File (TXT format)", data=translated_text, file_name=f"Translated_{filename}.txt")

    
            
                    
            #Can be used wherever a "file-like" object is accepted:
            #dataframe = pd.read_csv(uploaded_file)
            #st.write(dataframe)
            
    else:

        if "messages" not in st.session_state:
            st.session_state.messages = []


        res = [' ']
        # Chat input box
        text = st.chat_input("Type a text you want to translate")
        #st.write(f"{languages[target]}")
        if text:
            # Save user input to history
            st.session_state.messages.append({"role": "user", "content": text})

            # Display status while processing
            with st.status("Translating...", expanded=True) as status:
                # Simulated delay to mimic processing (replace with actual call)
                #time.sleep(2)  # Replace with the time your `generate` function takes
                if llm_model == "NMT":
                    contents = [text]
                    response = f"{generate_NMT(contents, languages[source], languages[target])[0]}"
                else:
                    response = f"{translate_text(text, languages[source], languages[target], llm_model, tone, domain, instruction, mandatory_translations)}"    # Replace with your `generate` function
                    st.write(f"{languages[target]}")
                    res.append(response)
                st.session_state.messages.append({"role": "assistant", "content": response})
                status.update(label="Translated", state="complete", expanded=True)

        
        for message in st.session_state.messages:
            if message["role"] == "user":
                st.chat_message("user").write(message["content"])
            else:
                st.chat_message("assistant").write(message["content"], key="copy_area")
        
        
        c1, c3 = st.columns([0.1, 0.4])
        with c1:
            if st.button("🗑️ Clear Chat"):
                st.session_state.messages = []

        with c3: 
            user_messages = [msg["content"] for msg in st.session_state.messages if msg["role"] == "user"]
            assistant_messages = [msg["content"] for msg in st.session_state.messages if msg["role"] == "assistant"]

            data1 = {"User": user_messages, "AI Response": assistant_messages}
            df_chat = pd.DataFrame(data1)

            df_chat.to_excel('Chats.xlsx', index=True)
            with open('Chats.xlsx', "rb") as template_file:
                template_byte = template_file.read()

            st.download_button(label="Download Chats",
                                data=template_byte,
                                file_name="Chats.xlsx",
                                mime='application/octet-stream')





